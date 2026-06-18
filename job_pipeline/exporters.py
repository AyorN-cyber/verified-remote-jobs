from __future__ import annotations

import csv
from dataclasses import asdict
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from .models import VerifiedProspect


FIELDS = [
    "status",
    "confidence_score",
    "company",
    "title",
    "candidate_fit",
    "apply_strategy",
    "linkedin_outreach_message",
    "email_followup_message",
    "materials_to_prepare",
    "company_research_notes",
    "outreach_strategy",
    "job_url",
    "apply_url",
    "company_url",
    "company_linkedin",
    "location",
    "salary",
    "source",
    "posted_at",
    "freshness_status",
    "freshness_evidence",
    "verification_summary",
    "official_source_status",
    "apply_link_status",
    "country_eligibility",
    "eligibility_evidence",
    "scam_risk",
    "rejection_reason",
    "evidence_links",
    "verified_at",
    "ai_notes",
]


def export_all(prospects: list[VerifiedProspect], out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    export_csv(prospects, out_dir / "verified_job_prospects.csv")
    export_xlsx(prospects, out_dir / "verified_job_prospects.xlsx")
    export_docx(prospects, out_dir / "verified_job_prospects_report.docx")


def export_csv(prospects: list[VerifiedProspect], path: Path) -> None:
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=FIELDS)
        writer.writeheader()
        for prospect in prospects:
            row = asdict(prospect)
            row["verified_at"] = prospect.verified_at.isoformat()
            writer.writerow({field: row.get(field, "") for field in FIELDS})


def export_xlsx(prospects: list[VerifiedProspect], path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Verified prospects"
    ws.append(FIELDS)
    for prospect in prospects:
        row = asdict(prospect)
        row["verified_at"] = prospect.verified_at.isoformat()
        ws.append([row.get(field, "") for field in FIELDS])

    status_fills = {
        "approved": PatternFill("solid", fgColor="C6EFCE"),
        "manual_review": PatternFill("solid", fgColor="FFF2CC"),
        "rejected": PatternFill("solid", fgColor="FFC7CE"),
    }
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(wrap_text=True, vertical="top")
    for row_idx in range(2, ws.max_row + 1):
        fill = status_fills.get(ws.cell(row_idx, 1).value, PatternFill("solid", fgColor="FFFFFF"))
        for col_idx in range(1, ws.max_column + 1):
            ws.cell(row_idx, col_idx).fill = fill
            ws.cell(row_idx, col_idx).alignment = Alignment(wrap_text=True, vertical="top")
    widths = {
        "status": 16,
        "confidence_score": 14,
        "company": 26,
        "title": 42,
        "candidate_fit": 46,
        "apply_strategy": 58,
        "linkedin_outreach_message": 58,
        "email_followup_message": 62,
        "materials_to_prepare": 44,
        "company_research_notes": 48,
        "outreach_strategy": 58,
        "job_url": 60,
        "apply_url": 60,
        "company_url": 36,
        "company_linkedin": 42,
        "rejection_reason": 52,
        "ai_notes": 52,
    }
    for idx, field in enumerate(FIELDS, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = widths.get(field, 24)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    wb.save(path)


def export_docx(prospects: list[VerifiedProspect], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    doc.add_heading("Verified Remote Job Prospects", 0)
    approved = [p for p in prospects if p.status == "approved"]
    manual = [p for p in prospects if p.status == "manual_review"]
    rejected = [p for p in prospects if p.status == "rejected"]
    doc.add_paragraph(
        "Only approved roles should be sent to the candidate. Manual-review roles need human checks. "
        "Rejected roles are retained for audit and should not be used for applications."
    )
    summary = doc.add_table(rows=1, cols=4)
    summary.style = "Table Grid"
    headers = ("Approved", "Manual Review", "Rejected", "Total")
    for idx, header in enumerate(headers):
        cell = summary.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    row = summary.add_row().cells
    row[0].text = str(len(approved))
    row[1].text = str(len(manual))
    row[2].text = str(len(rejected))
    row[3].text = str(len(prospects))

    if approved:
        doc.add_heading("Apply-Ready Prospects", 1)
        for item in approved:
            add_candidate_prospect(doc, item)

    if manual:
        doc.add_heading("Manual Review Prospects", 1)
        for item in manual:
            add_candidate_prospect(doc, item)

    if rejected:
        doc.add_heading("Rejected Audit Trail", 1)
        for item in rejected:
            add_audit_item(doc, item)
    doc.save(path)


def add_candidate_prospect(doc: Document, item: VerifiedProspect) -> None:
    doc.add_heading(f"{item.company} - {item.title}", 2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in (
        ("Status", f"{item.status} ({item.confidence_score}/100)"),
        ("Official Job Link", item.job_url),
        ("Apply Link", item.apply_url),
        ("Company Website", item.company_url),
        ("Company LinkedIn", item.company_linkedin),
        ("Location / Eligibility", f"{item.location} | {item.country_eligibility}: {item.eligibility_evidence}"),
        ("Salary", item.salary or "Not published"),
        ("Freshness", f"{item.freshness_status}: {item.freshness_evidence}"),
    ):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value or "Not verified"

    doc.add_heading("Why This Fits", 3)
    doc.add_paragraph(item.candidate_fit or "Not generated.")
    doc.add_heading("Apply Strategy", 3)
    doc.add_paragraph(item.apply_strategy or "Apply through official route only.")
    doc.add_heading("Company Research Notes", 3)
    doc.add_paragraph(item.company_research_notes or "Review website, product, customers, and support channels.")
    doc.add_heading("Materials To Prepare", 3)
    doc.add_paragraph(item.materials_to_prepare or "Tailored CV, cover note, LinkedIn profile, and certificate links.")
    doc.add_heading("LinkedIn Outreach Message", 3)
    doc.add_paragraph(item.linkedin_outreach_message or "Not generated.")
    doc.add_heading("Email Follow-Up Message", 3)
    doc.add_paragraph(item.email_followup_message or "Not generated.")
    doc.add_heading("Verification Notes", 3)
    doc.add_paragraph(item.verification_summary or "No summary.")
    doc.add_paragraph(f"Scam risk: {item.scam_risk or 'Not verified'}")


def add_audit_item(doc: Document, item: VerifiedProspect) -> None:
    doc.add_heading(f"{item.company} - {item.title}", 2)
    for label, value in (
        ("Source", item.source),
        ("Job URL", item.job_url),
        ("Freshness", f"{item.freshness_status}: {item.freshness_evidence}"),
        ("Eligibility", f"{item.country_eligibility}: {item.eligibility_evidence}"),
        ("Apply Link", item.apply_link_status),
        ("Scam Risk", item.scam_risk),
        ("Reason", item.rejection_reason),
    ):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        paragraph.add_run(value or "Not verified")
